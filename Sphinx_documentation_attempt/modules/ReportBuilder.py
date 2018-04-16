# -*- coding: utf-8 -*-
"""
Created on Wed Jan 10 10:41:54 2018

@author: h.gibb

Code to perform the construction of Calibration reports for Meters and Sources. 
"""
import docx
import time

class CalReport():
    def init(self, parent):
        
        """
        Initialisation requires the graphframe object in which the user enters report information. Text fields are read into local variables. 
        """

        self.name = parent.m_textCtrl188.GetValue()
        self.serial = parent.m_textCtrl189.GetValue()
        self.date = parent.m_textCtrl190.GetValue()
        self.client = parent.m_textCtrl191.GetValue()
        self.cond = parent.m_textCtrl196.GetValue()
        self.method =  parent.m_textCtrl192.GetValue()
        self.results = parent.m_textCtrl185.GetValue()
        self.parent = parent
        
      
        
    def BuildReport(self):
        
        """
        Report building produces a docx file that contains information from user fields.
        """
        
        if self.parent.m_checkBox1.GetValue():
            self.MeterCalculations()
        else:
            self.SourceCalculations()
            
        
        doc = docx.Document()
        doc.add_heading('Report on the Calibration of a ' + self.name, 1)
        doc.add_heading('Description', 3)
        doc.add_paragraph('A ' + self.name)
        
        doc.add_heading('Identification' , 3)        
        doc.add_paragraph('Serial: ' + self.serial)

        doc.add_heading('Client', 3)
        doc.add_paragraph(self.client)
        
        doc.add_heading('Date of Calibration', 3)
        doc.add_paragraph(self.date)
        
        doc.add_heading('Conditions of Calibration', 3)
        doc.add_paragraph(self.cond)
        
        doc.add_heading('Method', 3)
        doc.add_paragraph(self.method)
        
        doc.add_paragraph('The measurements were made with a delay of at least '+self.parent.m_textCtrl92.GetValue()+' seconds after making the connection between leads and setting the ranges.')
        doc.add_heading('DC Voltage Offset', 3)
        table = doc.add_table(rows=self.parent.m_grid41.GetNumberRows()+1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Instrument Range'
        hdr_cells[1].text = 'Instrument Readout'
        hdr_cells[2].text = 'Expanded Uncertainty'
        for i in range(0, 3):
            for j in range(0,self.parent.m_grid41.GetNumberRows()):
                table.rows[j+1].cells[i].text = str(self.parent.m_grid41.GetCellValue(j,i)) + 'V'
                print(str(self.parent.m_grid41.GetCellValue(j,i)))
                
                
        doc.add_heading('DC Voltage Gain', 3)       
        table2 = doc.add_table(rows=self.parent.m_grid42.GetNumberRows()+1, cols=5)
        hdr_cells2 = table2.rows[0].cells
        hdr_cells2[0].text = 'Instrument Range'
        hdr_cells2[1].text = 'Instrument Readout'
        hdr_cells2[2].text = 'Voltage +'
        hdr_cells2[3].text = 'Voltage -'
        hdr_cells2[4].text = 'Expanded Uncertainty'
        for i in range(0, 4):
            for j in range(0,self.parent.m_grid42.GetNumberRows()):
                table2.rows[j+1].cells[i].text = str(self.parent.m_grid42.GetCellValue(j,i)) + 'V'
                print(str(self.parent.m_grid42.GetCellValue(j,i)))

        doc.add_heading('DC Voltage Linearity', 3)
        table3 = doc.add_table(rows=self.parent.m_grid43.GetNumberRows()+1, cols=5)
        hdr_cells3 = table3.rows[0].cells
        hdr_cells3[0].text = 'Instrument Range'
        hdr_cells3[1].text = 'Instrument Readout'
        hdr_cells3[2].text = 'Voltage +'
        hdr_cells3[3].text = 'Voltage -'
        for i in range(0, 3):
            for j in range(0,self.parent.m_grid43.GetNumberRows()):
                table3.rows[j+1].cells[i].text = str(self.parent.m_grid43.GetCellValue(j,i)) + 'V'
                print(str(self.parent.m_grid43.GetCellValue(j,i)))

        
        
        doc.save('Calibration Report for '+self.name+time.strftime(" %a, %d %b %Y", time.localtime())+'.docx')
        
        
        
    def MeterCalculations(self):
        
        """
        Calculations for producing tables in a meter Calibration Report
        Calculate absolute values from ratios calculated in analysis (Yet to add)
        """

        self.col = range(0, self.parent.m_grid44.GetNumberCols()-1)
        
        self.labels = [self.col]
        self.expUnc = [self.col]
        self.ratio = [self.col]
        self.vRange = [self.col]
         
        for x in self.col:
            print(x)
            label = self.parent.m_grid44.GetCellValue(x, 0)
            ratio = self.parent.m_grid44.GetCellValue(x, 1)
            stDev = self.parent.m_grid44.GetCellValue(x, 2) 
            vRange = self.parent.m_grid44.GetCellValue(x, 4)
            if label[0][0] == 'm':
                self.labels[x] = label
                self.expUnc[x] = stDev #needs an operation
                self.ratio[x] = ratio 
                self.vRange[x] = vRange
                if vRange == 10:
                    self.abs[x] = 10*ratio
#                else:
#                    self.abs[x] = ratio*
#                self.abs[x] = 
        
        
    def SourceCalculations(self):

        """
        Calculations for producing tables in a source Calibration Report
        Calculate absolute values from ratios calculated in analysis(yet to add)
        """

        self.row = range(0, self.parent.m_grid44.GetNumberCols())
        
        self.labels = [self.row]
        self.expUnc = [self.row]
        self.ratio = [self.row]
        self.vRange = [self.row]
        
        for x in self.row:
            label = self.parent.m_grid44.GetCellValue(x, 0)
            ratio = self.parent.m_grid44.GetCellValue(x, 1)
            stDev = self.parent.m_grid44.GetCellValue(x, 2) 
            vRange = self.parent.m_grid44.GetCellValue(x, 4)
            if label[0][0] == 'x':
                self.labels[x] = label
                self.expUnc[x] = stDev #needs an operation
                self.ratio[x] = ratio #Not in final form for use in report.
                self.vRange[x] = vRange
                