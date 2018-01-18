# -*- coding: utf-8 -*-
"""
Created on Wed Jan 10 10:41:54 2018

@author: h.gibb

Code to perform the construction of Calibration reports for Meters and Sources. 
"""
import docx


class CalReport():
    def init(self, parent):
        """
        Initialisation requires the graphframe object in which the user enters report information. Text fields are read into local variables. 
        """
        self.name = parent.m_textCtrl188.GetValue()
        self.serial = parent.m_textCtrl189.GetValue()
        self.date = parent.m_textCtrl190.GetValue()
        self.client = parent.m_textCtrl191.GetValue()
        self.cond = parent.m_textCtrl196.GetValue()
        self.method =  parent.m_textCtrl192.GetValue()
        self.results = parent.m_textCtrl185.GetValue()
        self.parent = parent
      
        
    def BuildReport(self):
        """
        Report building produces a docx file that contains information from user fields. (adding in data from analysis)
        """
        if self.parent.m_checkBox1.GetValue():
            self.MeterCalculations()
        else:
            self.SourceCalculations()
            
        
        doc = docx.Document()
        doc.add_heading('Report on the Calibration of a ' + self.name, 1)
        doc.add_heading('Description', 3)
        doc.add_paragraph('A ' + self.name)
        
        doc.add_heading('Identification' , 3)        
        doc.add_paragraph('Serial: ' + self.serial)

        doc.add_heading('Client', 3)
        doc.add_paragraph(self.client)
        
        doc.add_heading('Date of Calibration', 3)
        doc.add_paragraph(self.date)
        
        doc.add_heading('Conditions of Calibration', 3)
        doc.add_paragraph(self.cond)
        
        doc.add_heading('Method', 3)
        doc.add_paragraph(self.method)
        
        doc.add_paragraph('')
        
        table = doc.add_table(rows=5, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Instrument Range'
        hdr_cells[1].text = 'Instrument Readout'
        hdr_cells[2].text = 'Expanded Uncertainty'
        doc.save('Calibration Report for '+self.name+'.docx')
        
    def MeterCalculations(self):
        """
        Calculations for producing tables in a meter Calibration Report
        Calculate absolute values from ratios calculated in analysis 
        """
        self.row = range(0, self.parent.m_grid44.GetNumberCols())
        
        self.labels = [self.row]
        self.expUnc = [self.row]
        self.ratio = [self.row]
        self.vRange = [self.row]
        
        for x in self.row:
            label = self.parent.m_grid44.GetCellValue(x, 0)
            ratio = self.parent.m_grid44.GetCellValue(x, 1)
            stDev = self.parent.m_grid44.GetCellValue(x, 2) 
            vRange = self.parent.m_grid44.GetCellValue(x, 4)
            if label[0][0] == 'm':
                self.labels[x] = label
                self.expUnc[x] = stDev #needs an operation
                self.ratio[x] = ratio 
                self.vRange[x] = vRange
#                self.abs[x] = 
                
                
    def SourceCalculations(self):
        """
        Calculations for producing tables in a source Calibration Report
        Calculate absolute values from ratios calculated in analysis
        """
        self.row = range(0, self.parent.m_grid44.GetNumberCols())
        
        self.labels = [self.row]
        self.expUnc = [self.row]
        self.ratio = [self.row]
        self.vRange = [self.row]
        
        for x in self.row:
            label = self.parent.m_grid44.GetCellValue(x, 0)
            ratio = self.parent.m_grid44.GetCellValue(x, 1)
            stDev = self.parent.m_grid44.GetCellValue(x, 2) 
            vRange = self.parent.m_grid44.GetCellValue(x, 4)
            if label[0][0] == 'x':
                self.labels[x] = label
                self.expUnc[x] = stDev #needs an operation
                self.ratio[x] = ratio #Not in final form for use in report.
                self.vRange[x] = vRange
                